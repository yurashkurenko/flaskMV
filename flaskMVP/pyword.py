from docx import Document
import os


def main():
    template_file_path = 'employment_agreement_template.docx'
    output_file_path = 'result.docx'

    variables = {
        "${EMPLOEE_NAME}": "Example Name",
        "${EMPLOEE_TITLE}": "Software Engineer",
        "${EMPLOEE_ID}": "302929393",
        "${EMPLOEE_ADDRESS}": "דרך השלום מנחם בגין דוגמא",
        "${EMPLOEE_PHONE}": "+972-5056000000",
        "${EMPLOEE_EMAIL}": "example@example.com",
        "${START_DATE}": "03 Jan, 2021",
        "${SALARY}": "10,000",
        "${SALARY_30}": "3,000",
        "${SALARY_70}": "7,000",
    }

    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


if __name__ == '__main__':
    main()